# os library for getting file path info and creating files
import os
# request library for getting url info and web scraping
import requests
# openpyxl is a Python library to read/write Excel
import openpyxl
# docx library for working with Word Document
from docx import Document
# Now Pt is used for font size in docx
from docx.shared import Pt
# BeautifulSoup will help us to get data from html code
from bs4 import BeautifulSoup
# Workbook will be used to enter data into Excel Sheets
from openpyxl import Workbook
# Access OpenPyXL columns using indices instead of letters
from openpyxl.utils import get_column_letter
# basename will get the current working directory path
from os.path import basename
# Counter will count the number of elements in a list
from collections import Counter

# base_url is the main link of our website
base_url = "https://www.ourcommons.ca"
# url will take us to the page where all members info is provided
url = "https://www.ourcommons.ca/Members/en/search"

document = Document()
# get the directory where our .py file is stored
pwd = os.path.dirname(__file__)
cwd = os.getcwd()
# If our current working directory is  different then give the path of the directory where .py file is stored
cd = os.chdir(pwd)

# html_page = requests.get(url)
# soup = BeautifulSoup(html_page.content, 'html.parser')
# warning = soup.find(
#     'div', class_="ce-mip-mp-picture-container")

# images = warning.findAll('img')
# image = images[0]
# source = image.attrs['src']
# link = f'{base_url}{source}'

# To fetch the data we used request command with GET method
response = requests.request(method='GET', url=url)
# if data is fetched properly it will return status code 200
if response.status_code == 200:
    # now beautifulsoup is a Python library for pulling data out of HTML and XML files
    soup = BeautifulSoup(response.text, features='lxml')
    # use find_all to get all div classes with provided class name
    data = soup.find_all('div', {'class': 'ce-mip-mp-tile-container'})
    # information is list in which we are going to store the data fetched in form of
    # dictionary where each dictionary is element of this list
    information = []
    # loop through the data list and fetch some more data
    for i in data:
        # now data for each member will be stored in info dictionary
        info = {}
        row = i.find_all('div', {'class': 'ce-mip-flex-tile'})
        # now Access the parent which have id and some text use string index to only store the id
        # using indexing Only get the last five digits which is id of member
        info["member_id"] = i.parent.find('div')['id'][-5:]
        # info["name"] will have name of the member
        info["name"] = i.find('div', {'class': 'ce-mip-mp-name'}).text
        # links will get the anchor tag
        links = i.find('a')
        # to fetch the url use 'get' after combining it with base url store it in info["links"]
        info["links"] = f"{base_url}{links.get('href')}"
        # here info about party is fetched
        info["party"] = i.find('div', {'class': 'ce-mip-mp-party'}).text
        # constituency
        info["constituency"] = i.find(
            'div', {'class': 'ce-mip-mp-constituency'}).text
        # province
        info["province"] = i.find('div', {'class': 'ce-mip-mp-province'}).text
        # url of image is stored in info["photo"]
        info["photo"] = f"{base_url}{i.find('img')['src']}"
        # Now add this all data stored in info dictionary to information list
        information.append(info)


# now we will write the data fetched into excel sheet

# Working with excel

# initialise wb as Workbook
wb = Workbook()
dest_filename = "report.xlsx"  # name by which the excel file is stored

# Make Sheet as active sheet
Sheet = wb.active
# To store the data 333 members we will use for loop
# Here we start loop from 1 as there is no 0th row and as range will loop till 332 we added 1 here
for i in range(1, (len(information)+1)):
    # Now i is a changing value so using string formatting it can be easily assigned with alphabets
    # we are using i-1 because dictionary starts its indexing from 0 and we are starting the loop from 1
    Sheet[f'A{i}'] = information[i-1]['member_id']
    Sheet[f'B{i}'] = information[i-1]['name']
    Sheet[f'C{i}'] = information[i-1]['party']
    Sheet[f'D{i}'] = information[i-1]['constituency']
    Sheet[f'E{i}'] = information[i-1]['province']
    Sheet[f'G{i}'] = information[i-1]['photo']
    Sheet[f'H{i}'] = information[i-1]['links']

# To change the height of each row use for loop
for row in range(Sheet.max_row+1):
    # the following code is for increasing the height of the row
    Sheet.row_dimensions[row].height = 175

# To change the width of each column use for loop
for col in range(Sheet.max_row):
    # We have to mention the column name to change its width
    Sheet.column_dimensions['A'].width = 10
    Sheet.column_dimensions['B'].width = 20
    Sheet.column_dimensions['C'].width = 20
    Sheet.column_dimensions['D'].width = 20
    Sheet.column_dimensions['E'].width = 20
    Sheet.column_dimensions['F'].width = 20
    Sheet.column_dimensions['G'].width = 90
    Sheet.column_dimensions['H'].width = 70

# To store the image link create a list
imgName = []

# Use the information list and loop through it
for info in information:
    # check if filename exist or not
    if os.path.isfile(basename(info["photo"])):
        imgName.append(info["photo"].split('/')[-1])
    else:
        # if it doesn't exist create image file using request.get() to download the file
        with open(basename(info["photo"]), "wb") as f:
            f.write(requests.get(info["photo"]).content)
        # now store the name of the file so that next time it doesn't need to download if it already exist
        imgName.append(info["photo"].split('/')[-1])

# to add image in excel use loop
for i in range(len(imgName)):
    # openpyxl is used to access the image file stored in imgName list
    img = openpyxl.drawing.image.Image(imgName[i])
    # Add the image in excel
    Sheet.add_image(img, f'F{i+1}')

# Save the Excel file
wb.save(dest_filename)


# Working with Document
document = Document()
# style some format
style = document.styles['Normal']
# font name
style.font.name = 'Cambria(Body)'
font = style.font
# font size in points
font.size = Pt(13)

# Add heading
document.add_heading('Province Wise Member Count', 0)
# provinces will store all 333 member's province
provinces = []
# use for loop to iterate through the information list
for i in range(len(information)):
    # access the province of each member and add it to provinces list
    provinces.append(information[i]["province"])

# Use counter to count number of each province in provinces and store that data in count_province
count_province = list(Counter(provinces).items())

# create table in document with 0 rows and 2 columns
table = document.add_table(rows=0, cols=2)
# add a header row
hdr1_cell = table.add_row().cells

# now the following code is to write the header in bold text (took some help from google)
header_1_1 = hdr1_cell[0].paragraphs[0].add_run(str('Province'))
header_1_1.bold = True

header_1_2 = hdr1_cell[1].paragraphs[0].add_run(str('Number of Members'))
header_1_2.bold = True

# for loop in count province to display number of members in each province
for i in count_province:
    # create a new row
    new_row = table.add_row().cells
    # in first row display name of province
    new_row[0].text = i[0]
    # in second row display number of members
    new_row[1].text = str(i[1])

# To add a page break
document.add_page_break()

# Again a heading on next page
document.add_heading('Party Wise Member Count', 0)
# parties will store all 333 member's party
parties = []
# use for loop to iterate through the information list
for i in range(len(information)):
    # access the party of each member and add it to parties list
    parties.append(information[i]["party"])

# Use counter to count number of each party in parties and store that data in count_province
count_party = list(Counter(parties).items())

# Create a table with 0 rows and 2 columns
table = document.add_table(rows=0, cols=2)
# add a header row
hdr2_cell = table.add_row().cells
# now the following code is to write the header in bold text (took some help from google)
header_2_1 = hdr2_cell[0].paragraphs[0].add_run(str('Political Party'))
header_2_1.bold = True

header_2_2 = hdr2_cell[1].paragraphs[0].add_run(str('Number of Members'))
header_2_2.bold = True

# for loop in count_party to display number of members in each party
for i in count_party:
    # create new row
    new_row = table.add_row().cells
    # in first row display name of each party
    new_row[0].text = i[0]
    # in second row display the number of members
    new_row[1].text = str(i[1])

# To add a page break
document.add_page_break()

# province_party will store
province_party = []
#
unique_provinces = Counter(provinces).keys()
# iterate through the unique_provinces and
for pro in unique_provinces:
    # initialise a dictionary party_dict and for each province store number of members from each party
    party_dict = {}
    # iterate through the parties list which have party of  each province
    for party in parties:
        # now initialize every key with value 0
        party_dict[party] = 0
    # now iterate through the information list
    for info in information:
        # check if province in unique_provinces matches province in information
        if pro == info["province"]:
            # if it matches add increase value of that party in pro(province in unique_provinces) province
            party_dict[info["party"]] += 1
    # append this party_dict to province_party
    province_party.append(party_dict)

# Give a heading in docx file
document.add_heading('Dominated Provinces', 0)
# now create a empty list in which wwe store all province names
ourUniqueProvinces = []
# using for loop iterate through provinces list and store the province if it is not in ourUniqueProvinces
for x in provinces:
    if x not in ourUniqueProvinces:
        ourUniqueProvinces.append(x)

# using for loop we will check which province has maximum member from which party
for i in range(len(ourUniqueProvinces)):
    # key will store the province with maximum members in a particular province
    key = max(province_party[i], key=lambda x: province_party[i][x])
    # to print the result we use document.add_paragraph()
    p = document.add_paragraph('In ')
    # to continue writing in same paragraph use add_run with previously assigned variable p
    p.add_run(f'{ourUniqueProvinces[i]}, {key} ').bold = True
    # To make the text bold use .bold = True
    p.add_run(f'is leading with ')
    p.add_run(f'{province_party[i][key]} ').bold = True
    p.add_run(f'Members of Parliaments.')

# Finally save the docx file and give it name report.docx
document.save("report.docx")
